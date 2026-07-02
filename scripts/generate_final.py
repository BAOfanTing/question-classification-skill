# -*- coding: utf-8 -*-
"""Final version: correct classification with 马哲, 毛思想+近代史, 中特+习思想."""
import os, re
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

CATEGORIES = [
    "刑法", "民法", "宪法", "行政法与诉讼法", "商法与经济法",
    "管理", "公文", "经济",
    "马哲",
    "毛思想+近代史",
    "中特+习思想",
    "人文", "时政", "地理科技"
]

FILES_DIR = "E:/KaoGong/JiangYi/到时候打印的/分类题目/"

def extract_answer(text):
    m = re.search(r'(?:答案[：:]\s*|【答案】\s*)([√×✅❌ABCDabcd、，,\s]+)', text)
    if m: return m.group(1).strip().rstrip(')）.】】')
    m = re.search(r'\*\*答案[：:](.*?)\*\*', text)
    if m: return m.group(1).strip()
    return ""

def determine_type(ans):
    if not ans: return "单选"
    if '√' in ans or '×' in ans: return "判断"
    cleaned = re.sub(r'[、，,\s]', '', ans)
    if len(cleaned) > 1 and all(c in 'ABCD' for c in cleaned.upper()): return "多选"
    return "单选"

def get_date_key(s):
    m = re.search(r'(\d{4})年(\d{1,2})月', s)
    return int(m.group(1)) * 100 + int(m.group(2)) if m else 0

# ===== TARGETED CLASSIFICATION =====
# Strategy: Only reclassify when very certain. Most questions stay in their original category.

def classify_meta_section(text):
    """
    Classify questions from the '中特（含哲学、马原、毛中特等内容）' meta-section.
    Split into: 马哲 / 毛思想+近代史 / 中特+习思想
    Priority: 马哲 > 毛思想+近代史 > 中特+习思想
    """
    # --- 马哲: Only clearly philosophical questions ---
    # These are questions that explicitly ask about philosophical concepts
    mz_strong = [
        "哲学基本问题", "唯物主义", "唯心主义",
        "古代朴素唯物主义", "近代形而上学唯物主义", "辩证唯物主义",
        "主观唯心主义", "客观唯心主义",
        "辩证法", "形而上学",
        "对立统一规律", "量变质变", "否定之否定",
        "矛盾的同一性和斗争性", "矛盾的普遍性和特殊性",
        "唯物辩证法的", "辩证法的规律",
        "认识论", "实践是认识", "真理是",
        "感性认识", "理性认识",
        "物质与意识", "思维和存在",
        "物质是不依赖于",
        "社会存在决定社会意识", "社会存在和社会意识",
        "经济基础与上层建筑",
        "生产力决定生产关系", "生产关系对生产力",
        "客观唯心", "主观唯心",
        "原子和虚空",
        "哲学范畴",
        "辩证唯物主义物质观",
        "物质概念",
        "从哲学上看", "哲学道理", "哲学原理",
        "包含的哲学", "体现的哲学",
        "物质世界的", "意识的本质",
        "实践是检验真理的唯一标准",
        "相对静止", "绝对运动",
        "人民群众是历史的",
        "社会历史性",
        "价值的主体性", "价值的客观性",
        "事物发展的趋势", "事物发展的规律",
        "真理都有",
        "哲学上", "哲学中",
        "属于辩证唯物主义",
        "唯物主义的观点",
        "主观能动性",
        "规律的客观性",
        "新旧事物",
        "具体问题具体分析",
        "一切从实际出发",
        "尊重客观规律",
        "整体与部分",
        "量变与质变",
        "矛盾的主要方面",
        "主要矛盾和次要矛盾",
        "社会基本矛盾",
        "唯物史观",
        "唯心史观",
    ]
    for kw in mz_strong:
        if kw in text:
            return "马哲"

    # --- 毛思想+近代史 ---
    mzds = [
        "毛泽东", "毛思想", "毛同志",
        "新民主主义革命", "新民主主义",
        "工农联盟", "根据地",
        "井冈山", "遵义会议",
        "八七会议", "中共七大",
        "洛川会议", "瓦窑堡",
        "三大纪律", "八项注意",
        "抗日战争", "解放战争",
        "长征", "红军",
        "半殖民地", "半封建",
        "反帝反封建",
        "三大改造", "社会主义改造",
        "辛亥革命", "五四运动",
        "中共一大", "中国共产党成立",
        "鸦片战争", "太平天国",
        "洋务运动", "戊戌变法",
        "义和团",
        "北洋军阀", "国共合作",
        "土地革命", "农村包围城市",
        "武装夺取政权",
        "党的建设", "统一战线", "武装斗争",
        "中国革命", "中国革命的",
        "近代中国",
        "百团大战", "平型关",
        "抗日民族统一战线",
        "延安整风", "七届二中",
        "抗美援朝",
        "社会主义革命",
        "旧民主主义",
        "民族独立", "人民解放",
        "无产阶级是中国",
        "分清敌友",
    ]
    for kw in mzds:
        if kw in text:
            return "毛思想+近代史"

    # --- Default: 中特+习思想 ---
    return "中特+习思想"


def classify_by_section(header_line):
    """Determine base category from the source section header."""
    header_map = {
        "刑法": "刑法",
        "民法": "民法",
        "宪法": "宪法",
        "商法": "商法与经济法",
        "经济法": "商法与经济法",
        "管理": "管理",
        "公文": "公文",
        "经济": "经济",
        "人文": "人文",
        "时政": "时政",
        "地理科技": "地理科技",
        "习思想": "中特+习思想",
    }
    for key, val in header_map.items():
        if key in header_line:
            return val
    return None


# ===== Parse All Files =====
all_data = OrderedDict()
for cat in CATEGORIES:
    all_data[cat] = []

for fname in sorted(os.listdir(FILES_DIR)):
    if not fname.endswith(".md") or fname.startswith("merge") or fname.startswith("generate"):
        continue
    fpath = os.path.join(FILES_DIR, fname)
    with open(fpath, encoding='utf-8') as f:
        content = f.read()

    content = re.sub(r'^# .+\n', '', content)
    sections = re.split(r'\n(?=## )', content)

    for sec in sections:
        if not sec.strip(): continue
        header_line = sec.split("\n")[0].strip()
        body = "\n".join(sec.split("\n")[1:])

        # Determine base category from section header
        is_meta_section = "中特" in header_line and ("哲学" in header_line or "毛" in header_line or "马原" in header_line or "中特" in header_line)
        base_cat = classify_by_section(header_line)

        if base_cat is None and not is_meta_section:
            continue

        blocks = re.split(r'\n(?=\d*\s*\.?\s*\*{0,2}\(?\(?\d{4}年\d{1,2}月)', body)

        for block in blocks:
            block = block.strip()
            if not block or len(block) < 10: continue
            if block.startswith("###") or block == "---": continue

            date_m = re.search(r'\((\d{4}年\d{1,2}月\s*[^)]+)\)', block)
            if not date_m: continue
            date_str = date_m.group(1).strip()
            date_key = get_date_key(date_str)

            answer = extract_answer(block)
            if not answer: continue
            q_type = determine_type(answer)

            q_text = block
            q_text = re.sub(r'\*\*答案[：:].*?\*\*', '', q_text)
            q_text = re.sub(r'(?:答案[：:]\s*|【答案】\s*)[√×✅❌ABCDabcdefg、，,\s]+', '', q_text)
            q_text = re.sub(r'\*\*', '', q_text).strip()
            q_text = re.sub(r'^\d+\s*[.、]\s*', '', q_text).strip()
            q_text = re.sub(r'\n\s*-{3,}', '', q_text)
            q_text = re.sub(r'^-{3,}\s*', '', q_text).strip()

            # Determine target category
            if is_meta_section:
                target_cat = classify_meta_section(q_text)
            else:
                target_cat = base_cat

            # Reclassify ALL questions for 行政法与诉讼法 (originally mixed in 宪法 etc.)
            admin_kw = [
                "行政", "行政处罚", "行政许可", "行政强制", "行政复议",
                "行政赔偿", "行政诉讼", "治安管理处罚",
                "行政监察", "行政拘留", "行政处分",
                "国家赔偿", "治安管理",
            ]
            for kw in admin_kw:
                if kw in q_text and target_cat != "行政法与诉讼法":
                    # Check it's really 行政法, not just 行政 in general context
                    # Words like 行政机关, 行政行为 are strong signals
                    if re.search(r'行政(机关|行为|处罚|许可|强制|复议|诉讼|赔偿|监察|拘留|处分|法规|管理|执法)', q_text):
                        target_cat = "行政法与诉讼法"
                        break

            if target_cat:
                all_data[target_cat].append({
                    'date_key': date_key,
                    'date_str': date_str,
                    'q_text': q_text,
                    'answer': answer,
                    'q_type': q_type
                })

# Sort
type_order = {"单选": 0, "多选": 1, "判断": 2}
for cat in all_data:
    all_data[cat].sort(key=lambda x: (x['date_key'], type_order.get(x['q_type'], 3)))

# ===== Generate DOCX =====
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title
title = doc.add_heading('综合基础知识真题分类汇总', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("来源：浙江宁波/杭州事业单位《综合基础知识》真题14套（2021年6月—2023年8月）")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# TOC
h = doc.add_heading('目  录', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
run = p.add_run()
run._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w'))))
run2 = p.add_run()
run2._element.append(parse_xml(r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'.format(nsdecls('w'))))
run3 = p.add_run()
run3._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w'))))
run4 = p.add_run("（在Word中右键此处→更新域→更新整个目录）")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run4.font.italic = True
run5 = p.add_run()
run5._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w'))))
doc.add_page_break()

# Content
total_q = 0

for cat_name in CATEGORIES:
    items = all_data[cat_name]
    if not items: continue

    h2 = doc.add_heading(cat_name, level=1)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    by_type = {"单选": [], "多选": [], "判断": []}
    for item in items:
        by_type[item['q_type']].append(item)

    all_answers = []
    q_counter = 1

    for qtype in ["单选", "多选", "判断"]:
        type_items = by_type[qtype]
        if not type_items: continue
        doc.add_heading(qtype, level=2)

        for item in type_items:
            q_text = item['q_text']
            q_text = re.sub(r'^\(?\s*\d{4}年\d{1,2}月\s*[^)]*\)?\s*', '', q_text).strip()
            q_text = re.sub(r'^[.、\s]+', '', q_text).strip()
            q_text = re.sub(r'^\d+\s*[.、]\s*', '', q_text).strip()

            lines = q_text.split('\n')
            # Split inline options like "A.xxx  B.xxx  C.xxx  D.xxx" onto separate lines
            new_lines = []
            for line in lines:
                # Check if line has multiple options separated by 2+ spaces
                if re.search(r'[A-D][.、)]', line) and re.search(r'\s{2,}[A-D][.、)]', line):
                    # Split by option pattern with preceding spaces
                    parts = re.split(r'(?=\s{2,}[A-D][.、)])', line)
                    for part in parts:
                        part = part.strip()
                        if part:
                            new_lines.append(part)
                else:
                    new_lines.append(line)
            lines = new_lines

            p = doc.add_paragraph(f"{q_counter}.({item['date_str']}) {lines[0]}")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)

            for opt_line in lines[1:]:
                opt_line = opt_line.strip()
                if opt_line and (re.match(r'^[A-D①②③④⑤⑥⑦⑧⑨⑩]', opt_line) or re.match(r'^\d+[.、]', opt_line)):
                    p = doc.add_paragraph(opt_line)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.left_indent = Cm(0.5)

            all_answers.append((q_counter, item['answer']))
            q_counter += 1

    if all_answers:
        doc.add_paragraph()
        doc.add_heading('答案', level=2)
        buf = []
        ans_groups = []
        for num, ans in all_answers:
            buf.append(f"{num}.{ans}")
            if len(buf) >= 12:
                ans_groups.append("  ".join(buf))
                buf = []
        if buf: ans_groups.append("  ".join(buf))
        for line in ans_groups:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

    total_q += len(items)
    doc.add_page_break()

# Page numbers
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w'))))
    run2 = p.add_run()
    run2._element.append(parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w'))))
    run3 = p.add_run()
    run3._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w'))))
    run4 = p.add_run('1')
    run4.font.size = Pt(10)
    run5 = p.add_run()
    run5._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w'))))

output_path = "E:/KaoGong/JiangYi/到时候打印的/分类题目/汇总/综合基础知识真题分类汇总.docx"
doc.save(output_path)

# Stats
print("分类统计：")
print("=" * 50)
for cat in CATEGORIES:
    items = all_data[cat]
    if items:
        bt = {"单选": 0, "多选": 0, "判断": 0}
        for i in items: bt[i['q_type']] += 1
        detail = "  ".join([f"{k}{v}" for k, v in bt.items() if v > 0])
        print(f"  {cat:12s}  {len(items):3d}题  ({detail})")
    else:
        print(f"  {cat:12s}  0题")

print(f"\n总计: {total_q}题")
print(f"输出: {output_path}")
